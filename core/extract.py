"""Turns an incoming resume of any supported format into a flat list of Lines.

Each Line keeps the few visual signals the parser needs to tell a heading from
body text: whether Word/PDF marked it as a list item, whether it is bold, and
its font size.
"""

import re
from dataclasses import dataclass
from pathlib import Path

BULLET_CHARS = "\u2022\u25cf\u25aa\u25e6\u2023\u2043\u00b7\u2219\uf0b7\uf0a7\uf0d8\u2013\u2014-*>"
_BULLET_RE = re.compile(rf"^\s*[{re.escape(BULLET_CHARS)}]+\s+")


@dataclass
class Line:
    text: str
    is_bullet: bool = False
    bold: bool = False
    size: float | None = None


class UnsupportedFormat(Exception):
    pass


def extract(path: str) -> list[Line]:
    suffix = Path(path).suffix.lower()
    if suffix == ".docx":
        lines = _from_docx(path)
    elif suffix == ".pdf":
        lines = _from_pdf(path)
    elif suffix in (".txt", ".md", ".text"):
        lines = _from_text(Path(path).read_text(encoding="utf-8", errors="replace"))
    elif suffix == ".rtf":
        lines = _from_text(_strip_rtf(Path(path).read_text(encoding="utf-8", errors="replace")))
    elif suffix == ".doc":
        lines = _from_legacy_doc(path)
    else:
        raise UnsupportedFormat(
            f"Cannot read '{suffix}' files.\n\n"
            "Supported: .pdf, .docx, .doc, .rtf, .txt\n"
            "Tip: open the file in Word and 'Save As' .docx, then try again."
        )
    return _normalise(lines)


def _normalise(lines: list[Line]) -> list[Line]:
    out = []
    for ln in lines:
        text = ln.text.replace("\xa0", " ").replace("\t", " ")
        text = re.sub(r"[ ]{2,}", "  ", text).strip()
        if not text:
            continue
        if _BULLET_RE.match(text):
            ln.is_bullet = True
            text = _BULLET_RE.sub("", text)
        # A lone bullet glyph on its own line carries no content.
        if not text.strip(BULLET_CHARS + " "):
            continue
        ln.text = text.strip()
        out.append(ln)
    return out


def _from_docx(path: str) -> list[Line]:
    import docx
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(path)
    lines: list[Line] = []

    def emit(par: Paragraph) -> None:
        text = par.text
        if not text.strip():
            return
        pPr = par._p.find(qn("w:pPr"))
        is_bullet = pPr is not None and pPr.find(qn("w:numPr")) is not None
        if not is_bullet and "list" in (par.style.name or "").lower():
            is_bullet = True
        meaningful = [r for r in par.runs if r.text.strip()]
        bold = bool(meaningful) and all(r.bold for r in meaningful)
        size = next((r.font.size.pt for r in meaningful if r.font.size), None)
        lines.append(Line(text, is_bullet, bold, size))

    def walk(parent) -> None:
        for child in parent.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                emit(Paragraph(child, doc))
            elif tag == "tbl":
                # Two-column resume layouts are common; read cells in order.
                for row in Table(child, doc).rows:
                    for cell in row.cells:
                        for par in cell.paragraphs:
                            emit(par)

    # Some agencies put the candidate's name in the page header rather than the
    # body, so read it first — it is the strongest name candidate when present.
    for section in doc.sections:
        for par in section.header.paragraphs:
            emit(par)

    walk(doc.element.body)
    return lines


def _from_pdf(path: str) -> list[Line]:
    import pdfplumber

    lines: list[Line] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(extra_attrs=["fontname", "size"])
            if not words:
                continue
            rows: dict[int, list[dict]] = {}
            for w in words:
                # Bucket by baseline; 3pt absorbs sub-pixel jitter within a line.
                key = round(w["top"] / 3)
                rows.setdefault(key, []).append(w)
            for key in sorted(rows):
                row = sorted(rows[key], key=lambda w: w["x0"])
                text = " ".join(w["text"] for w in row)
                fonts = [w.get("fontname", "") or "" for w in row]
                bold = all(("bold" in f.lower() or "black" in f.lower()) for f in fonts)
                sizes = [w.get("size") for w in row if w.get("size")]
                lines.append(Line(text, False, bold, max(sizes) if sizes else None))
    return lines


def _from_text(raw: str) -> list[Line]:
    return [Line(t) for t in raw.splitlines()]


def _strip_rtf(raw: str) -> str:
    raw = re.sub(r"\\par[d]?", "\n", raw)
    raw = re.sub(r"\{\\\*?[^{}]*\}", "", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d*\s?", "", raw)
    return raw.replace("{", "").replace("}", "")


def _from_legacy_doc(path: str) -> list[Line]:
    """Legacy binary .doc has no pure-Python reader; drive Word if installed."""
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        raise UnsupportedFormat(
            "Old-style .doc files need Microsoft Word installed.\n\n"
            "Easiest fix: open the file in Word and 'Save As' .docx, then try again."
        )

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(path).resolve()), ReadOnly=True)
        text = doc.Content.Text
        doc.Close(False)
        return _from_text(text.replace("\r", "\n"))
    except Exception as exc:
        raise UnsupportedFormat(
            f"Could not read the .doc file ({exc}).\n\n"
            "Open it in Word and 'Save As' .docx, then try again."
        )
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()
