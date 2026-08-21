"""Renders ResumeData into the ShimentoX client format.

All formatting (page size, styles, bullet numbering, header logo) is inherited
from templates/shimentox.docx, which is the client's own sample with its body
emptied. This module only lays out content.
"""

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Emu, Inches, Pt, Twips

from .model import ResumeData

FONT = "Times New Roman"
SIZE = Pt(10)
BULLET_NUM_ID = 4  # the Symbol-bullet list defined in the donor template
BULLET_INDENT_TWIPS = 144

DISPLAY_NAME = "ShimentoX"


def template_path() -> Path:
    from .paths import resource

    return resource("templates/shimentox.docx")


def render(data: ResumeData, out_path: Path) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template_path(), out_path)

    doc = Document(str(out_path))
    _set_header(doc, data.name)

    if data.summary:
        _heading(doc, "Summary:")
        for item in data.summary:
            _run(_para(doc, bullet=True), item)
        _blank(doc)

    if data.skills:
        _heading(doc, "Technical Skills:")
        for group in data.skills:
            p = _para(doc, bullet=True)
            if group.category:
                _run(p, f"{group.category}:", bold=True)
                _run(p, f" {group.values}")
            else:
                _run(p, group.values)
        _blank(doc)

    if data.experience:
        _heading(doc, "Professional Experience:")
        for job in data.experience:
            p = _para(doc, right_tab=True, keep_with_next=True)
            _run(p, job.company, bold=True)
            if job.dates:
                _run(p, "\t", bold=True)
                _run(p, job.dates, bold=True, italic=True)
            if job.title:
                _run(_para(doc), job.title, bold=True, italic=True)
            if job.project:
                pp = _para(doc)
                _run(pp, "Project: ", bold=True, italic=True)
                _run(pp, job.project, italic=True)
            if job.bullets:
                _run(_para(doc), "Responsibilities:", bold=True, italic=True)
                for bullet in job.bullets:
                    _run(_para(doc, bullet=True), bullet)
            _blank(doc)

    if data.certifications:
        _heading(doc, "Certifications:")
        for cert in data.certifications:
            _run(_para(doc, bullet=True), cert)
        _blank(doc)

    if data.education:
        _heading(doc, "Education:")
        for edu in data.education:
            p = _para(doc, right_tab=True)
            label = ", ".join(x for x in (edu.degree, edu.institution) if x)
            _run(p, label)
            if edu.year:
                _run(p, "\t")
                _run(p, edu.year, bold=True, italic=True)

    doc.save(str(out_path))
    return out_path


def _set_header(doc: Document, name: str) -> None:
    """Build a stable two-column header with the logo flush to the right margin."""
    from .paths import resource

    section = doc.sections[0]
    header = section.header
    for child in list(header._element):
        header._element.remove(child)
    usable_width = section.page_width - section.left_margin - section.right_margin
    table = header.add_table(rows=1, cols=2, width=usable_width)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table._tbl.tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
        '<w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        '</w:tblBorders>'
    ))

    left_width = int(usable_width * 0.65)
    right_width = usable_width - left_width
    left, right = table.rows[0].cells
    for cell, width in ((left, left_width), (right, right_width)):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(parse_xml(
            f'<w:tcMar {nsdecls("w")}><w:top w:w="0" w:type="dxa"/>'
            '<w:left w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/>'
            '<w:right w:w="0" w:type="dxa"/></w:tcMar>'
        ))

    name_paragraph = left.paragraphs[0]
    name_paragraph.paragraph_format.space_after = Pt(0)
    _run(name_paragraph, name, bold=True).font.size = Pt(14)

    logo_paragraph = right.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_paragraph.paragraph_format.space_after = Pt(0)
    logo_paragraph.add_run().add_picture(
        str(resource("assets/shimento_logo.png")), width=Inches(2.08)
    )


def _para(
    doc: Document,
    bullet: bool = False,
    right_tab: bool = False,
    keep_with_next: bool = False,
):
    p = doc.add_paragraph(style="No Spacing")
    # Full justification creates visibly stretched headings and short bullets.
    # The client sample uses a clean left-aligned resume body.
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.keep_with_next = keep_with_next
    p.paragraph_format.widow_control = True
    if bullet:
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/>'
            f'<w:numId w:val="{BULLET_NUM_ID}"/></w:numPr>'
        ))
        p.paragraph_format.left_indent = Twips(BULLET_INDENT_TWIPS)
    if right_tab:
        section = doc.sections[-1]
        right_edge = section.page_width - section.left_margin - section.right_margin
        p.paragraph_format.tab_stops.add_tab_stop(
            Emu(right_edge), WD_TAB_ALIGNMENT.RIGHT
        )
    return p


def _blank(doc: Document) -> None:
    _para(doc)


def _heading(doc: Document, text: str) -> None:
    _run(_para(doc, keep_with_next=True), text, bold=True)


def _run(p, text: str, bold: bool = False, italic: bool = False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = SIZE
    r.bold = bold
    r.italic = italic
    # python-docx sets ascii/hAnsi only; complex-script needs setting by hand or
    # Word falls back to the theme font for any non-Latin character.
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:cs"), FONT)
    return r
