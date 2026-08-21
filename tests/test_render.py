import shutil
import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.model import Education, Job, ResumeData, SkillGroup
from core.export import to_pdf
from core.render_shimentox import render


class RenderTests(unittest.TestCase):
    def test_template_embeds_high_resolution_aspect_correct_logo(self):
        template = Path(__file__).parents[1] / "templates" / "shimentox.docx"
        with zipfile.ZipFile(template) as archive:
            image = archive.read("word/media/image1.png")
            header = ET.fromstring(archive.read("word/header1.xml"))

        width = int.from_bytes(image[16:20], "big")
        height = int.from_bytes(image[20:24], "big")
        namespace = {
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        }
        extent = header.find(".//wp:extent", namespace)
        self.assertIsNotNone(extent)
        self.assertGreaterEqual(width, 1000)
        self.assertAlmostEqual(
            int(extent.attrib["cx"]) / int(extent.attrib["cy"]),
            width / height,
            places=2,
        )

    def test_render_preserves_content_and_client_layout(self):
        data = ResumeData(
            name="Jane Doe",
            summary=["Delivers reliable software."],
            skills=[SkillGroup("Languages", "Python, SQL")],
            experience=[Job("Example Ltd", "2022 - Present", "Engineer", "Atlas",
                            ["Improved processing time."])],
            certifications=["Cloud Practitioner"],
            education=[Education("B.Tech", "Example University", "2022")],
        )
        with tempfile.TemporaryDirectory() as directory:
            output = render(data, Path(directory) / "resume.docx")
            document = Document(output)
            body = "\n".join(p.text for p in document.paragraphs)
            header = " ".join(
                cell.text
                for table in document.sections[0].header.tables
                for row in table.rows
                for cell in row.cells
            )

        self.assertIn("Jane Doe", header)
        for expected in ("Summary:", "Technical Skills:", "Professional Experience:",
                         "Certifications:", "Education:", "Example Ltd", "2022 - Present"):
            self.assertIn(expected, body)
        self.assertTrue(all(p.paragraph_format.space_after.pt == 0
                            for p in document.paragraphs if p.paragraph_format.space_after))

    def test_repeated_render_keeps_all_supported_resume_data(self):
        for iteration in range(1, 21):
            data = ResumeData(
                name=f"Candidate {iteration}",
                summary=[f"Summary {iteration}-{index}" for index in range(1, 6)],
                skills=[SkillGroup(f"Skill group {index}", f"Value {iteration}-{index}")
                        for index in range(1, 5)],
                experience=[Job(f"Company {index}", f"202{index} - Present",
                                f"Role {index}", f"Project {index}",
                                [f"Responsibility {iteration}-{index}-{bullet}"
                                 for bullet in range(1, 4)])
                            for index in range(1, 4)],
                certifications=[f"Certification {iteration}-{index}" for index in range(1, 3)],
                education=[Education(f"Degree {index}", f"University {index}", f"202{index}")
                           for index in range(1, 3)],
            )
            with tempfile.TemporaryDirectory() as directory:
                output = render(data, Path(directory) / "resume.docx")
                document = Document(output)
                text = "\n".join(paragraph.text for paragraph in document.paragraphs)

            expected = [*data.summary, *(group.values for group in data.skills),
                        *(job.company for job in data.experience),
                        *(bullet for job in data.experience for bullet in job.bullets),
                        *data.certifications, *(education.degree for education in data.education)]
            for value in expected:
                self.assertIn(value, text, f"iteration {iteration} lost {value!r}")

    def test_logo_is_right_aligned_in_every_generated_header(self):
        with tempfile.TemporaryDirectory() as directory:
            output = render(ResumeData(name="Header Check"), Path(directory) / "resume.docx")
            document = Document(output)
            header = document.sections[0].header
            self.assertEqual(len(header.tables), 1)
            right_paragraph = header.tables[0].cell(0, 1).paragraphs[0]
            self.assertEqual(right_paragraph.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertEqual(len(right_paragraph._p.xpath(".//a:blip")), 1)

    @unittest.skipUnless(shutil.which("soffice"), "LibreOffice is not installed")
    def test_end_to_end_docx_and_pdf_export(self):
        data = ResumeData(name="UAT Candidate", summary=["End-to-end export check."])
        with tempfile.TemporaryDirectory() as directory:
            docx = render(data, Path(directory) / "uat-resume.docx")
            pdf = to_pdf(docx)
            self.assertGreater(docx.stat().st_size, 0)
            self.assertGreater(pdf.stat().st_size, 0)
            self.assertEqual(pdf.read_bytes()[:4], b"%PDF")


if __name__ == "__main__":
    unittest.main()
